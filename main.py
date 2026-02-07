import os
import re
import textwrap
from sqlalchemy.future import select
from sqlalchemy import delete
from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, Request, Form, UploadFile, File, HTTPException, Depends
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy.sql import not_
from sqlalchemy.orm import selectinload
from sqlalchemy.orm.attributes import flag_modified
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

from database import engine, Base, get_db, AsyncSessionLocal
from models import Room, Message, PushSubscription

from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

import uuid

import asyncio

import redis.asyncio as redis
import json
from contextlib import asynccontextmanager
from google import genai

from pywebpush import webpush, WebPushException

from utils.summarizer import generate_summary, ask_gemini_simple, analyze_discussion_progress, get_facilitation_from_gemini

from utils.client import get_gemini_client
import textwrap
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
redis_client = None

VAPID_PRIVATE_KEY = os.getenv("VAPID_PRIVATE_KEY")
VAPID_PUBLIC_KEY = os.getenv("VAPID_PUBLIC_KEY")

VAPID_CLAIMS = {
    "sub": "mailto:admin@example.com"
}

@asynccontextmanager
async def lifespan(app: FastAPI):
    # アプリケーション起動時に実行
    print("アプリケーションを起動します...")
    # データベースのテーブルを作成
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    
    # Redisへの接続
    global redis_client
    redis_url = os.getenv("REDIS_URL", "redis://localhost:6379/0")
    try:
        redis_client = redis.from_url(redis_url, encoding="utf-8", decode_responses=True)
        await redis_client.ping()
        print("Redisに正常に接続しました。")
    except Exception as e:
        print(f"Redisへの接続に失敗しました: {e}")
        redis_client = None

    yield  # ここでアプリケーションが実行される

    # アプリケーション終了時に実行
    print("アプリケーションを終了します...")
    if redis_client:
        await redis_client.close()
        print("Redisとの接続を閉じました。")

app = FastAPI(lifespan=lifespan)

app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# 1. Service Worker配信
@app.get("/sw.js")
async def service_worker():
    return FileResponse("sw.js", media_type="application/javascript")

# 2. 通知購読用のデータモデルとAPI
class SubscriptionSchema(BaseModel):
    endpoint: str
    keys: dict
    username: str
    room_id: str

@app.post("/subscribe")
async def subscribe(sub: SubscriptionSchema, db: AsyncSession = Depends(get_db)):
    result = await db.execute(
        select(PushSubscription).filter_by(endpoint=sub.endpoint, room_id=sub.room_id)
    )
    existing_sub = result.scalars().first()

    if existing_sub:
        if existing_sub.username != sub.username:
            existing_sub.username = sub.username
            await db.commit()
    else:
        new_sub = PushSubscription(
            room_id=sub.room_id,
            username=sub.username,
            endpoint=sub.endpoint,
            p256dh=sub.keys['p256dh'],
            auth=sub.keys['auth']
        )
        db.add(new_sub)
        await db.commit()
    return {"message": "Subscribed successfully"}

# 3. 通知送信のヘルパー関数
async def send_push_notification(subscriptions: list, message_content: str, sender_name: str, exclude_users: set):
    try:
        if not VAPID_PRIVATE_KEY or not VAPID_PUBLIC_KEY:
            return

        # DBアクセスは削除し、渡されたリストを使用する
        for sub in subscriptions:
            if sub.username == sender_name or sub.username in exclude_users:
                continue
            
            payload = json.dumps({
                "title": f"{sender_name}さんからのメッセージ",
                "body": message_content[:100],
                "url": f"/room/{sub.room_id}?username={sub.username}" # sub.room_idを使用
            })

            try:
                await asyncio.to_thread(
                    webpush,
                    subscription_info={
                        "endpoint": sub.endpoint,
                        "keys": {"p256dh": sub.p256dh, "auth": sub.auth}
                    },
                    data=payload,
                    vapid_private_key=VAPID_PRIVATE_KEY,
                    vapid_claims=VAPID_CLAIMS
                )
            except Exception as e:
                print(f"Notification Error for {sub.username}: {e}")

    except Exception as e:
        print(f"Error in send_push_notification: {e}")

UPLOAD_DIR = "static/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

EXCEL_DIR = "static/excels"
os.makedirs(EXCEL_DIR, exist_ok=True)

FONT_PATH = "fonts/ipaexg.ttf"
if not os.path.exists("fonts"):
    os.makedirs("fonts")

try:
    pdfmetrics.registerFont(TTFont('IPAexGothic', FONT_PATH))
    print(f"Font 'IPAexGothic' registered from {FONT_PATH}")
except Exception as e:
    print(f"Warning: Could not register font from {FONT_PATH}. Error: {e}")

# ... (index, create_room, join_room, room 関数の部分は変更なし) ...
@app.get("/", response_class=HTMLResponse)
async def index(request: Request, db: AsyncSession = Depends(get_db), error: str = None):
    result = await db.execute(select(Room))
    all_rooms = result.scalars().all()
    rooms_info = [
        {
            "room_id": room.room_id,
            "topic": room.topic,
            "status": room.status
        }
        for room in all_rooms
    ]
    return templates.TemplateResponse("index.html", {
        "request": request,
        "error": error,
        "rooms": rooms_info,
    })

@app.post("/create")
async def create_room(username: str = Form(...), topic: str = Form(...), db: AsyncSession = Depends(get_db)):
    room_id = str(uuid.uuid4())[:8]
    new_room = Room(
        room_id=room_id,
        topic=topic,
        analytics={"users": {}}, # analyticsの初期構造を作成
        proposals_data=[]
    )
    db.add(new_room)
    await db.commit()
    print(f"Room created: {room_id} - topic: {topic}")
    return JSONResponse(content={
        "room_id": room_id,
        "username": username,
        "url": f"/room/{room_id}?username={username}"
    })

@app.post("/join")
async def join_room(request: Request, username: str = Form(...), room_id: str = Form(...), db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Room).filter_by(room_id=room_id))
    room = result.scalars().first()
    if not room:
        return RedirectResponse(f"/?error=無効なルームIDです", status_code=303)
    return RedirectResponse(f"/room/{room_id}?username={username}", status_code=303)


@app.get("/room/{room_id}", response_class=HTMLResponse)
async def room(request: Request, room_id: str, username: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Room).filter_by(room_id=room_id))
    room_obj = result.scalars().first()
    if not room_obj:
        return RedirectResponse(f"/?error=存在しないルームです", status_code=303)
    
    return templates.TemplateResponse("chat.html", {
        "request": request,
        "username": username,
        "room_id": room_obj.room_id,
        "topic": room_obj.topic,
        "vapid_public_key": VAPID_PUBLIC_KEY 
    })

@app.websocket("/ws/{room_id}/{username}")
async def websocket_endpoint(websocket: WebSocket, room_id: str, username: str):
    # 【変更点】引数から db: AsyncSession = Depends(get_db) を削除しました

    if not redis_client:
        await websocket.accept()
        await websocket.send_json({"type": "system_message", "content": "サーバーエラー: データベース接続に問題があります。"})
        await websocket.close()
        return

    # 1. ルーム存在確認とAnalytics初期化 (必要な時だけDBを開く)
    async with AsyncSessionLocal() as db:
        result = await db.execute(
            select(Room).options(selectinload(Room.messages)).filter_by(room_id=room_id)
        )
        room_obj = result.scalars().first()
        
        if not room_obj:
            await websocket.accept()
            await websocket.send_json({"type": "system_message", "content": "エラー: 指定されたルームは存在しません。"})
            await websocket.close()
            return

        try:
                result_lock = await db.execute(select(Room).filter_by(room_id=room_id).with_for_update())
                room_obj_for_analytics = result_lock.scalars().first()
                
                if room_obj_for_analytics:
                    analytics = room_obj_for_analytics.analytics
                    if "users" not in analytics:
                        analytics["users"] = {}
                    
                    if username not in analytics["users"]:
                        analytics["users"][username] = {
                            "posts": 0, "stances": {}, 
                            "reactions_given": {"agree": 0, "partial": 0, "disagree": 0},
                            "reactions_received": {"agree": 0, "partial": 0, "disagree": 0},
                            "note_edits": 0, "facilitator_uses": 0, "proposal_form_edits": 0
                        }
                        flag_modified(room_obj_for_analytics, "analytics")

                        await db.commit()
        except Exception as e:
            print(f"Analytics への接続ユーザー登録でエラー: {e}")

    await websocket.accept()
    
    room_channel = f"room:{room_id}"
    participants_key = f"participants:{room_id}"

    # --- Readerタスク ---
    async def reader(pubsub, event):
        try:
            await pubsub.subscribe(room_channel)
            event.set()
            while True:
                message = await pubsub.get_message(ignore_subscribe_messages=True, timeout=None)
                if message and message.get("data"):
                    try:
                        await websocket.send_json(json.loads(message["data"]))
                    except RuntimeError:
                        break
        except asyncio.CancelledError:
            pass
        except Exception as e:
            print(f"Readerタスクでエラーが発生 (Room: {room_id}, User: {username}): {e}")
        finally:
            if pubsub:
                await pubsub.unsubscribe(room_channel)

    # --- Writerタスク ---
    async def writer():
        try:
            while True:
                data = await websocket.receive_json()
                
                # 【重要】メッセージ受信時、その都度DB接続を開くように変更
                async with AsyncSessionLocal() as db:
                    async with db.begin():
                        # ロックを取得して競合を防ぐ
                        result = await db.execute(select(Room).filter_by(room_id=room_id).with_for_update())
                        room_obj_writer = result.scalars().first()
                        if not room_obj_writer:
                            continue

                        message_type = data.get("type")

                        if message_type == "message":
                            stance = data["stance"].strip()
                            content = data["content"]

                            analytics = room_obj_writer.analytics
                            user_analytics = analytics.get("users", {}).get(username)
                            
                            if user_analytics:
                                user_analytics["posts"] += 1
                                user_analytics["stances"][stance] = user_analytics.get(stance, 0) + 1
                                flag_modified(room_obj_writer, "analytics")

                            parent_message_dict = None
                            if data.get("reply_to_id"):
                                parent_res = await db.execute(select(Message).filter_by(message_id=data["reply_to_id"]))
                                parent_message = parent_res.scalars().first()
                                if parent_message:
                                    parent_message_dict = {
                                        "id": parent_message.message_id, 
                                        "username": parent_message.username, 
                                        "content": parent_message.content
                                    }

                            new_message = Message(
                                room_id=room_id,
                                username=username,
                                content=content,
                                stance=stance,
                                file_url=data.get("file_url"),
                                original_filename=data.get("original_filename"),
                                gemini_file_ref=data.get("gemini_file_ref"), 
                                reply_to_id=data.get("reply_to_id")
                            )
                            db.add(new_message)
                            await db.flush()

                            message_to_send = new_message.to_dict(parent_message_dict)
                            
                            # 通知処理用データの取得
                            sub_result = await db.execute(select(PushSubscription).filter_by(room_id=room_id))
                            subscriptions = sub_result.scalars().all()
                            subscriptions_data = [
                                PushSubscription(
                                    endpoint=sub.endpoint, p256dh=sub.p256dh, auth=sub.auth, 
                                    username=sub.username, room_id=sub.room_id
                                ) for sub in subscriptions
                            ]
                            current_online_users = await redis_client.smembers(participants_key)
                            
                            asyncio.create_task(
                                send_push_notification(subscriptions_data, content, username, set(current_online_users))
                            )

                            if stance == "Geminiへの質問":
                                await redis_client.publish(room_channel, json.dumps({"type": "message", **message_to_send}))
                                
                                files_to_ask = [new_message.gemini_file_ref] if new_message.gemini_file_ref else []
                                gemini_answer = await asyncio.to_thread(
                                    ask_gemini_simple, question=content, files=files_to_ask
                                )
                                
                                answer_message_obj = Message(
                                    room_id=room_id,
                                    username="Gemini",
                                    content=gemini_answer,
                                    stance="Geminiからの回答"
                                )
                                db.add(answer_message_obj)
                                await redis_client.publish(room_channel, json.dumps({"type": "gemini_response", **answer_message_obj.to_dict()}))
                            else:
                                await redis_client.publish(room_channel, json.dumps({"type": "message", **message_to_send}))

                        elif message_type == "reaction":
                            message_id = data.get("message_id")
                            reaction_type = data.get("reaction")
                            res = await db.execute(select(Message).filter_by(message_id=message_id))
                            target_message = res.scalars().first()
                            
                            if target_message:
                                previous_reaction = None
                                for r_type, users in target_message.reactions.items():
                                    if username in users:
                                        previous_reaction = r_type
                                        break
                                already_reacted = username in target_message.reactions.get(reaction_type, [])
                                
                                new_reactions = dict(target_message.reactions)
                                for r_type, users in new_reactions.items():
                                    if username in users:
                                        users.remove(username)
                                
                                if not already_reacted:
                                    if reaction_type not in new_reactions:
                                        new_reactions[reaction_type] = []
                                    new_reactions[reaction_type].append(username)
                                
                                target_message.reactions = new_reactions
                                flag_modified(target_message, "reactions")

                                analytics = room_obj_writer.analytics
                                if username in analytics.get("users", {}):
                                    author_username = target_message.username
                                    if previous_reaction:
                                        analytics["users"][username]["reactions_given"][previous_reaction] -= 1
                                        if author_username and author_username in analytics["users"]:
                                            analytics["users"][author_username]["reactions_received"][previous_reaction] -= 1
                                    if not already_reacted:
                                        analytics["users"][username]["reactions_given"][reaction_type] += 1
                                        if author_username and author_username in analytics["users"]:
                                            analytics["users"][author_username]["reactions_received"][reaction_type] += 1
                                    flag_modified(room_obj_writer, "analytics")
                                
                                reaction_update_data = {"type": "reaction_update", "message_id": message_id, "reactions": {k: len(v) for k, v in target_message.reactions.items()}}
                                await redis_client.publish(room_channel, json.dumps(reaction_update_data))

                        elif message_type == "delete_message":
                            message_id_to_delete = data.get("message_id")
                            res = await db.execute(select(Message).filter_by(message_id=message_id_to_delete))
                            message_to_delete = res.scalars().first()

                            if message_to_delete and message_to_delete.username == username:
                                analytics = room_obj_writer.analytics
                                if username in analytics.get("users", {}):
                                    user_analytics = analytics["users"][username]
                                    user_analytics["posts"] -= 1
                                    if message_to_delete.stance in user_analytics.get("stances", {}):
                                        user_analytics["stances"][message_to_delete.stance] -= 1
                                    flag_modified(room_obj_writer, "analytics")

                                await db.delete(message_to_delete)
                                await redis_client.publish(room_channel, json.dumps({
                                    "type": "message_deleted", 
                                    "message_id": message_id_to_delete
                                }))

                        elif message_type == "resolve_proposal":
                            message_id_to_resolve = data.get("message_id")
                            res = await db.execute(select(Message).filter_by(message_id=message_id_to_resolve))
                            message_to_resolve = res.scalars().first()
            
                            if message_to_resolve and message_to_resolve.stance == "提案":
                                message_to_resolve.is_resolved = True
                                flag_modified(message_to_resolve, "is_resolved")
                                await redis_client.publish(room_channel, json.dumps({
                                    "type": "proposal_resolved",
                                    "message_id": message_id_to_resolve
                                }))

                        elif message_type == "note_update":
                            if room_obj_writer.status != "終了":
                                content = data.get("content", "")
                                room_obj_writer.shared_note = content
                                analytics = room_obj_writer.analytics
                                if username in analytics.get("users", {}):
                                    analytics["users"][username]["note_edits"] += 1
                                flag_modified(room_obj_writer, "analytics")
                                note_update_data = {"type": "note_update", "content": content, "sender": username}
                                await redis_client.publish(room_channel, json.dumps(note_update_data))

                        elif message_type == "proposal_form_update":
                            if room_obj_writer.status != "終了":
                                proposals_list = data.get("proposals", [])
                                room_obj_writer.proposals_data = proposals_list
                                analytics = room_obj_writer.analytics
                                if username in analytics.get("users", {}):
                                    analytics["users"][username]["proposal_form_edits"] = analytics["users"][username].get("proposal_form_edits", 0) + 1
                                flag_modified(room_obj_writer, "analytics")
                                form_update_data = {
                                    "type": "proposal_form_update", 
                                    "proposals": proposals_list, 
                                    "sender": username
                                }
                                await redis_client.publish(room_channel, json.dumps(form_update_data))

                        elif message_type == "finish":
                            room_obj_writer.status = "終了"
                            system_message = {"type": "system_message", "content": "議事録を作成中です。しばらくお待ちください..."}
                            await redis_client.publish(room_channel, json.dumps(system_message))

                            res = await db.execute(
                                select(Message)
                                .filter(Message.room_id == room_id, Message.stance != "summary") 
                                .order_by(Message.created_at)
                            )
                            messages_from_db = res.scalars().all()
                            
                            chat_messages_dict = [msg.to_dict() for msg in messages_from_db]
                            files_for_summary = [msg.gemini_file_ref for msg in messages_from_db if msg.gemini_file_ref]
                            note_content = room_obj_writer.shared_note
                            proposals_data = room_obj_writer.proposals_data or []
                            topic = room_obj_writer.topic

                            summary_content = await asyncio.to_thread(
                                generate_summary, 
                                chat_messages_dict, 
                                topic, 
                                files=files_for_summary, 
                                note_content=note_content,
                                proposals_data=proposals_data
                            )
                            
                            participants = list(room_obj_writer.analytics.get("users", {}).keys())

                            excel_filename = f"meeting_minutes_{room_id}.xlsx"
                            excel_path = os.path.join(EXCEL_DIR, excel_filename)
                            excel_url = f"/excels/{excel_filename}"
                            
                            await asyncio.to_thread(
                                create_meeting_minutes_excel, 
                                messages_from_db, 
                                topic, 
                                participants, 
                                excel_path
                            )
                            
                            summary_data_dict = {
                                "content": summary_content,
                                "excel_url": excel_url  # キー名を pdf_url から excel_url に変更
                            }
                            
                            res_summary = await db.execute(select(Message).filter_by(room_id=room_id, stance="summary"))
                            summary_message_obj = res_summary.scalars().first()

                            if summary_message_obj:
                                summary_message_obj.content = json.dumps(summary_data_dict)
                                flag_modified(summary_message_obj, "content")
                            else:
                                summary_message_obj = Message(
                                    room_id=room_id,
                                    username="System", 
                                    content=json.dumps(summary_data_dict),
                                    stance="summary"
                                )
                                db.add(summary_message_obj)
                            
                            publish_data = {"type": "summary", **summary_data_dict}
                            await redis_client.publish(room_channel, json.dumps(publish_data))

        except WebSocketDisconnect:
            pass
        except Exception as e:
            print(f"Writerタスクでエラーが発生 (Room: {room_id}, User: {username}): {e}")

    pubsub = redis_client.pubsub()
    subscribe_event = asyncio.Event()
    reader_task = asyncio.create_task(reader(pubsub, subscribe_event))
    writer_task = asyncio.create_task(writer())
    await subscribe_event.wait()
    
    await redis_client.sadd(participants_key, username)
    current_participants = list(await redis_client.smembers(participants_key))
    update_message = {"type": "participant_update", "users": current_participants}
    await redis_client.publish(room_channel, json.dumps(update_message))

    # --- 履歴読み込み処理 ---
    try:
        # 【重要】ここでも必要な時だけDB接続を開く
        async with AsyncSessionLocal() as db:
            res_room = await db.execute(select(Room).filter_by(room_id=room_id))
            current_room_obj = res_room.scalars().first()
            if not current_room_obj:
                await websocket.close(code=1008, reason="Room not found during history load")
                return

            res_history = await db.execute(
                select(Message)
                .filter_by(room_id=room_id)
                .order_by(Message.created_at)
            )
            history_messages = res_history.scalars().all()
            message_dict_for_replies = {msg.message_id: msg for msg in history_messages}

            for message in history_messages:
                if message.stance == "summary":
                    summary_data = json.loads(message.content)
                    await websocket.send_json({"type": "summary", **summary_data})
                    continue
                
                parent_message_dict = None
                if message.reply_to_id:
                    parent_message = message_dict_for_replies.get(message.reply_to_id)
                    if parent_message:
                        parent_message_dict = {
                            "id": parent_message.message_id, 
                            "username": parent_message.username, 
                            "content": parent_message.content
                        }
                await websocket.send_json({"type": "history", **message.to_dict(parent_message_dict)})
            
            await websocket.send_json({"type": "note_initial_state", "content": current_room_obj.shared_note})
            await websocket.send_json({
                "type": "proposal_form_initial_state", 
                "proposals": current_room_obj.proposals_data if current_room_obj.proposals_data else []
            })
    except RuntimeError:
        print(f"History send failed for {username}: client disconnected.")
        writer_task.cancel()
        reader_task.cancel()

    done, pending = await asyncio.wait(
        [reader_task, writer_task],
        return_when=asyncio.FIRST_COMPLETED
    )
    for task in pending:
        task.cancel()
    
    try:
        if redis_client:
            try:
                await redis_client.srem(participants_key, username)
                remaining_participants = list(await redis_client.smembers(participants_key))
                if remaining_participants:
                     update_message = {"type": "participant_update", "users": remaining_participants}
                     await redis_client.publish(room_channel, json.dumps(update_message))
            except Exception as e:
                print(f"Redis cleanup error: {e}")
        
        if pubsub:
            await pubsub.close()
        print(f"User {username} disconnected from room {room_id}. Cleaned up resources.")
    except Exception:
        pass

class FacilitatePayload(BaseModel):
    username: str

class ProgressCheckPayload(BaseModel):
    username: str

@app.post("/check_progress/{room_id}")
async def check_progress(room_id: str, payload: ProgressCheckPayload, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Room).filter_by(room_id=room_id))
    room_obj = result.scalars().first()

    if not room_obj:
        return JSONResponse(content={"error": "ルームが見つかりません。"}, status_code=404)
    
    res_msg = await db.execute(
        select(Message)
        .filter(Message.room_id == room_id, Message.stance != "summary")
        .order_by(Message.created_at)
    )
    messages_from_db = res_msg.scalars().all()
    
    # MissingGreenletエラーを回避するため、commitの前に属性にアクセスする
    note_content = room_obj.shared_note
    
    try:
        username = payload.username
        
        # 分析データを更新
        analytics = room_obj.analytics
        if username in analytics.get("users", {}):
            # "progress_check_uses" というキーで回数をカウントアップ
            analytics["users"][username]["progress_check_uses"] = analytics["users"][username].get("progress_check_uses", 0) + 1
            flag_modified(room_obj, "analytics")
            await db.commit() # ここでコミット
        else:
            # ユーザーが analytics に見つからない場合（通常は発生しないはず）
            await db.rollback()
            
    except Exception as e:
        print(f"Error during progress check analytics update: {e}")
        await db.rollback()

    # --- 3. 読み取ったデータを使ってAIに問い合わせ ---
    chat_messages_dict = [msg.to_dict() for msg in messages_from_db]
    if not chat_messages_dict and not note_content:
        return JSONResponse(content={"progress": "まだ議論が開始されていません。"})

    files_for_analysis = [msg['gemini_file_ref'] for msg in chat_messages_dict if msg.get('gemini_file_ref')]

    progress_summary = await asyncio.to_thread(
        analyze_discussion_progress, chat_messages_dict, room_obj.topic, files=files_for_analysis, note_content=note_content, proposals_data=room_obj.proposals_data or []
    )
    return JSONResponse(content={"progress": progress_summary})

@app.post("/upload_file/")
async def upload_file_endpoint(file: UploadFile = File(...)):
    try:
        # 新しい作法でクライアントを取得
        client = get_gemini_client()
    except ValueError as e:
        return JSONResponse(content={"message": str(e)}, status_code=500)

    # Gemini APIがサポートするMIMEタイプのリスト
    SUPPORTED_MIME_TYPES = [
        "application/pdf", "image/png", "image/jpeg", "text/plain", "text/html",
        "text/css", "text/javascript", "application/json", "text/markdown", "text/x-python",
    ]

    if file.content_type not in SUPPORTED_MIME_TYPES:
        error_message = f"サポートされていないファイル形式です: {file.filename} ({file.content_type})。PDF, TXT, 画像ファイルなどを利用してください。"
        return JSONResponse(content={"message": error_message}, status_code=400)

    file_extension = os.path.splitext(file.filename)[1]
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    file_location = os.path.join(UPLOAD_DIR, unique_filename)

    try:
        # ファイルをサーバーに一時保存
        with open(file_location, "wb") as f:
            f.write(await file.read())
        
        # フロントエンドが参照するためのURL
        file_url = f"/static/uploads/{unique_filename}"

        # 新しいSDKのファイルアップロード方法
        uploaded_file = client.files.upload(file=file_location)

        print(f"Completed upload. File name: {uploaded_file.name}")

        # フロントエンドに返す情報をJSON形式で作成
        return JSONResponse(content={
            "file_url": file_url,
            "original_filename": file.filename,
            "gemini_file_ref": uploaded_file.name
        })
    except Exception as e:
        print(f"!!! 詳細なエラー内容: {e}")
        print(f"!!! エラーの型: {type(e)}")
        return JSONResponse(content={"message": f"ファイルのアップロード中にエラーが発生しました: {e}"}, status_code=500)

@app.get("/excels/{filename}")
async def get_excel(filename: str):
    file_path = os.path.join(EXCEL_DIR, filename)
    if os.path.exists(file_path):
        return FileResponse(
            file_path, 
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
            filename=filename
        )
    return JSONResponse(content={"message": "ファイルが見つかりません。"}, status_code=404)

def create_meeting_minutes_pdf(summary_text: str, pdf_path: str):
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    font_name = 'IPAexGothic' if 'IPAexGothic' in pdfmetrics.getRegisteredFontNames() else 'Helvetica'
    if font_name == 'Helvetica':
         print("Warning: IPAexGothic font not available, using default font for PDF.")
         
    normal_style = styles['Normal']
    normal_style.fontName = font_name
    normal_style.fontSize = 10
    normal_style.leading = 14

    title_style = styles['h1']
    title_style.fontName = font_name
    
    story = [Paragraph("会議議事録", title_style), Spacer(1, 0.2 * 28)]

    for line in summary_text.split('\n'):
        p_text = line.replace(' ', '&nbsp;').replace('\n', '<br/>')
        story.append(Paragraph(p_text, normal_style))
    
    try:
        doc.build(story)
        print(f"PDF generated: {pdf_path}")
    except Exception as e:
        print(f"Failed to generate PDF for {pdf_path}: {e}")

@app.post("/delete_room")
async def delete_room(room_id: str = Form(...), password: str = Form(...), db: AsyncSession = Depends(get_db)):
    # パスワードを環境変数から取得するように修正することを推奨
    delete_password = os.getenv("DELETE_PASSWORD", "nagai")
    if password != delete_password:
        return RedirectResponse(f"/?error=パスワードが違います。", status_code=303)

    result = await db.execute(select(Room).filter_by(room_id=room_id))
    room_to_delete = result.scalars().first()

    if room_to_delete:
        # 2. 【追加】Gemini上のファイルを削除する処理
        try:
            # このルームのメッセージで、Geminiファイル参照を持っているものを検索
            msg_result = await db.execute(
                select(Message)
                .filter_by(room_id=room_id)
                .where(Message.gemini_file_ref.is_not(None))
            )
            messages_with_files = msg_result.scalars().all()
            
            if messages_with_files:
                client = get_gemini_client()
                print(f"Deleting {len(messages_with_files)} files from Gemini...")
                
                for msg in messages_with_files:
                    try:
                        # Gemini APIを使ってファイルを削除
                        # msg.gemini_file_ref は "files/xxxx..." という形式の名前
                        client.files.delete(name=msg.gemini_file_ref)
                        print(f"Deleted Gemini file: {msg.gemini_file_ref}")
                    except Exception as e:
                        print(f"Failed to delete file {msg.gemini_file_ref}: {e}")

        except Exception as e:
            # ファイル削除に失敗しても、ルーム削除自体は続行する
            print(f"Error during Gemini file cleanup: {e}")

        # 3. DBからルームを削除（カスケード設定によりメッセージも消えます）
        await db.delete(room_to_delete)
        await db.commit()
        print(f"Room deleted: {room_id}")
        return RedirectResponse(f"/?", status_code=303)
    
    return RedirectResponse(f"/?error=削除対象のルームが見つかりません。", status_code=303)

@app.post("/facilitate/{room_id}")
async def facilitate_discussion(room_id: str, payload: FacilitatePayload, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Room).filter_by(room_id=room_id))
    room_obj = result.scalars().first()

    if not room_obj:
        raise HTTPException(status_code=404, detail="Room not found")

    try:
        username = payload.username
        
        # 分析データを更新
        analytics = room_obj.analytics
        if username in analytics.get("users", {}):
            analytics["users"][username]["facilitator_uses"] += 1
            flag_modified(room_obj, "analytics")

        res = await db.execute(
            select(Message)
            .filter(Message.room_id == room_id, Message.stance != "summary")
            .order_by(Message.created_at)
        )
        messages_from_db = res.scalars().all()
        chat_messages_dict = [msg.to_dict() for msg in messages_from_db]
        
        note_content = room_obj.shared_note
        facilitation_text = await asyncio.to_thread(
            get_facilitation_from_gemini, chat_messages_dict, room_obj.topic, note_content, proposals_data=room_obj.proposals_data or []
        )

        # AIの発言としてメッセージを作成し、DBに追加
        ai_message_obj = Message(
            room_id=room_id,
            username="Gemini（AIファシリテーター）",
            content=facilitation_text,
            stance="ファシリテーション"
        )
        db.add(ai_message_obj)
        await db.commit()

        # 全員にブロードキャスト
        if redis_client:
            room_channel = f"room:{room_id}"
            await redis_client.publish(room_channel, json.dumps({"type": "message", **ai_message_obj.to_dict()}))
        return JSONResponse(content={"status": "success", "message": facilitation_text})
    except Exception as e:
        print(f"Error during facilitation: {e}")
        raise HTTPException(status_code=500, detail=f"AIファシリテーション中にサーバーエラーが発生しました: {e}")
        
@app.get("/dev-tools", response_class=HTMLResponse)
async def dev_tools_page(request: Request, password: str = None):
    dev_password = os.getenv("DEV_PASSWORD", "nagai")

    if password != dev_password:
        return RedirectResponse(f"/?error=パスワードが違います。", status_code=303)
        
    return templates.TemplateResponse("dev_tools.html", {"request": request})

@app.get("/api/analytics")
async def get_analytics_data(db: AsyncSession = Depends(get_db)):
    try:
        result = await db.execute(select(Room))
        all_rooms = result.scalars().all()

    # DBから取得した analytics データを基に、フロントエンドで扱いやすい形に整形
        analytics_from_db = {room.room_id: room.analytics for room in all_rooms if room.analytics}
    
        by_room = {}
        overall = {
            "posts": 0, "stances": {}, "reactions_given": {"agree": 0, "partial": 0, "disagree": 0},
            "reactions_received": {"agree": 0, "partial": 0, "disagree": 0},
            "note_edits": 0, "facilitator_uses": 0, "participants": set()
        }

        for room_id, r_data in analytics_from_db.items():
            room_summary = {
                "posts": 0, "stances": {}, "reactions_given": {"agree": 0, "partial": 0, "disagree": 0},
                "reactions_received": {"agree": 0, "partial": 0, "disagree": 0},
                "note_edits": 0, "facilitator_uses": 0, "participants": set()
            }
            for user, u_data in r_data.get("users", {}).items():
                room_summary["posts"] += u_data.get("posts", 0)
                overall["posts"] += u_data.get("posts", 0)
                room_summary["note_edits"] += u_data.get("note_edits", 0)
                overall["note_edits"] += u_data.get("note_edits", 0)
                room_summary["facilitator_uses"] += u_data.get("facilitator_uses", 0)
                overall["facilitator_uses"] += u_data.get("facilitator_uses", 0)

                room_summary["proposal_form_edits"] = room_summary.get("proposal_form_edits", 0) + u_data.get("proposal_form_edits", 0)
                overall["proposal_form_edits"] = overall.get("proposal_form_edits", 0) + u_data.get("proposal_form_edits", 0)

                room_summary["progress_check_uses"] = room_summary.get("progress_check_uses", 0) + u_data.get("progress_check_uses", 0)
                overall["progress_check_uses"] = overall.get("progress_check_uses", 0) + u_data.get("progress_check_uses", 0)

                room_summary["participants"].add(user)
                overall["participants"].add(user)

                for stance, count in u_data.get("stances", {}).items():
                    room_summary["stances"][stance] = room_summary["stances"].get(stance, 0) + count
                    overall["stances"][stance] = overall["stances"].get(stance, 0) + count
                for r_type, count in u_data.get("reactions_given", {}).items():
                    room_summary["reactions_given"][r_type] = room_summary["reactions_given"].get(r_type, 0) + count
                    overall["reactions_given"][r_type] = overall["reactions_given"].get(r_type, 0) + count
                for r_type, count in u_data.get("reactions_received", {}).items():
                    room_summary["reactions_received"][r_type] = room_summary["reactions_received"].get(r_type, 0) + count
                    overall["reactions_received"][r_type] = overall["reactions_received"].get(r_type, 0) + count
        
            room_summary["participants"] = len(room_summary["participants"])
            by_room[room_id] = room_summary

        overall["participants"] = len(overall["participants"])
    
    # 最終的なレスポンスを作成
        response_data = {
            "by_room_by_user": analytics_from_db,
            "by_room": by_room,
            "overall": overall
        }
        return JSONResponse(content=response_data)
    except Exception as e:
        print(f"Error in get_analytics_data: {e}")
        raise HTTPException(status_code=500, detail="分析データの集計中にサーバーエラーが発生しました。")

class WordDownloadPayload(BaseModel):
    topic: str
    proposals: list

def set_cell_border(cell, **kwargs):
    """
    テーブルのセルに枠線を設定するヘルパー関数
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_east_asia_font(run, font_name='Meiryo'):
    """
    Runオブジェクトに対して日本語（東アジア）フォントを設定するヘルパー関数
    """
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    # w:rFonts 要素を取得または作成
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    # 東アジアフォントを設定
    rFonts.set(qn('w:eastAsia'), font_name)

@app.post("/download_proposals_word")
async def download_proposals_word(payload: WordDownloadPayload):
    try:
        # ドキュメントの作成
        doc = Document()
        
        # 標準スタイルのフォント設定
        style = doc.styles['Normal']
        style.font.name = 'Meiryo'
        # スタイルの場合は element 属性でアクセス可能
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')
        style.font.size = Pt(10.5)

        # --- タイトル ---
        # add_heading は Paragraph を返すが、中身のテキストは Run で管理する方が安全
        title = doc.add_heading(level=0)
        # 既存のテキストをクリアして Run を追加し直すことでフォントを確実に適用
        title.clear() 
        run = title.add_run(f"気候市民会議 提言案一覧")
        run.font.color.rgb = RGBColor(0, 0, 0)
        set_east_asia_font(run, 'Meiryo')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- サブタイトル（議題） ---
        subtitle = doc.add_paragraph()
        run = subtitle.add_run(f"議題: {payload.topic}")
        run.font.size = Pt(14)
        run.bold = True
        set_east_asia_font(run, 'Meiryo')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("") # スペーサー

        # 質問項目のマッピング
        labels = {
            "q1": "【What】 提案内容",
            "q2": "【Why】 なぜ大切か",
            "q6": "【How】 実施手法",
            "q3": "【When】 実施時期",
            "q4": "【Where】 実施場所",
            "q5": "【Who】 実施主体・対象",
            "q7": "思考法"
        }

        # --- 各提案のループ ---
        for idx, prop in enumerate(payload.proposals, 1):
            # 提案ヘッダー
            h = doc.add_heading(level=1)
            run = h.add_run(f"提案 {idx}")
            set_east_asia_font(run, 'Meiryo')
            
            # テーブル作成 (2列: 項目名, 内容)
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.allow_autofit = False
            
            # 各項目の行を追加
            order = ["q1", "q2", "q6", "q3", "q4", "q5", "q7"]
            
            for key in order:
                row = table.add_row()
                cell_label = row.cells[0]
                cell_content = row.cells[1]
                
                # ラベル列の設定
                cell_label.width = Cm(6)
                p_label = cell_label.paragraphs[0]
                run_label = p_label.add_run(labels[key])
                run_label.bold = True
                run_label.font.size = Pt(10)
                set_east_asia_font(run_label, 'Meiryo')

                # 内容列の設定
                content_text = prop.get(key, "")
                
                # Q7（思考法）の特別な処理
                if key == "q7":
                    if content_text == "forecast":
                        content_text = "フォアキャスティング"
                    elif content_text == "backcast":
                        content_text = "バックキャスティング"
                    else:
                        content_text = "未選択"
                
                final_text = content_text if content_text else "（未記入）"
                p_content = cell_content.paragraphs[0]
                run_content = p_content.add_run(final_text)
                set_east_asia_font(run_content, 'Meiryo')

            doc.add_paragraph("") # 提案間のスペース
            
            # 最後の提案以外は改ページを入れる
            if idx < len(payload.proposals):
                doc.add_page_break()

        # --- ファイル書き出し ---
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"proposals_{uuid.uuid4().hex[:8]}.docx"
        
        # ダウンロード用レスポンス
        return StreamingResponse(
            buffer, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        print(f"Word generation error: {e}")
        # スタックトレースを含めて詳細を出力するとデバッグしやすい
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Wordファイルの作成に失敗しました: {str(e)}")
    
    # main.py の末尾付近に追加

@app.get("/debug/reset_push_subscriptions")
async def reset_push_subscriptions(db: AsyncSession = Depends(get_db)):
    try:
        # PushSubscription テーブルの中身を全て削除
        await db.execute(delete(PushSubscription))
        await db.commit()
        return JSONResponse(content={
            "status": "success", 
            "message": "通知の購読データを全て削除しました。ブラウザで再度「通知ON」の設定を行ってください。"
        })
    except Exception as e:
        await db.rollback()
        return JSONResponse(content={"status": "error", "message": str(e)}, status_code=500)
    
def create_meeting_minutes_excel(messages, topic, participants, file_path):
    wb = Workbook()
    ws = wb.active
    ws.title = "議事録"

    # ヘッダー情報
    ws["A1"] = "議題"
    ws["B1"] = topic
    ws["A2"] = "参加者"
    ws["B2"] = ", ".join(participants) if participants else "なし"

    # スタイル設定
    bold_font = Font(bold=True, name="Meiryo UI")
    normal_font = Font(name="Meiryo UI")
    center_align = Alignment(horizontal="center", vertical="center")
    top_left_align = Alignment(horizontal="left", vertical="top", wrap_text=True)
    top_center_align = Alignment(horizontal="center", vertical="top")
    thin_border = Border(left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin"))

    # メタデータスタイル
    for cell in ["A1", "B1", "A2", "B2"]:
        ws[cell].font = normal_font
    ws["A1"].font = bold_font
    ws["A2"].font = bold_font

    # テーブルヘッダー
    headers = ["No.", "発言者", "発言の種類", "発言内容", "リアクション数"]
    header_row = 4
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col_idx, value=header)
        cell.font = bold_font
        cell.alignment = center_align
        cell.border = thin_border

    # データ行
    for i, msg in enumerate(messages, 1):
        row_idx = header_row + i
        reaction_sum = sum(len(users) for users in (msg.reactions or {}).values())

        row_data = [i, msg.username, msg.stance, msg.content, reaction_sum]

        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = normal_font
            cell.border = thin_border
            
            if col_idx == 1 or col_idx == 5: cell.alignment = top_center_align # No, リアクション
            elif col_idx == 4: cell.alignment = top_left_align # 内容
            else: cell.alignment = top_center_align

    # 列幅調整
    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 15
    ws.column_dimensions["D"].width = 60
    ws.column_dimensions["E"].width = 15

    wb.save(file_path)